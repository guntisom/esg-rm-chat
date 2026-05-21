import io
import json
import math
import os
import re

import requests
from anthropic import Anthropic
from dotenv import load_dotenv
from tavily import TavilyClient

load_dotenv()


def _get_secret(key: str) -> str:
    try:
        import streamlit as st
        return st.secrets[key]
    except Exception:
        return os.getenv(key, "")


tavily = TavilyClient(api_key=_get_secret("TAVILY_API_KEY"))
claude = Anthropic(api_key=_get_secret("ANTHROPIC_API_KEY"))


def _fmt(val):
    if val is None:
        return None
    try:
        f = float(val)
        if math.isnan(f):
            return None
        return round(f / 1_000_000, 1)
    except Exception:
        return None


def fetch_financials(ticker: str) -> dict:
    ticker_bk = f"{ticker.upper()}.BK"
    api_key = _get_secret("FMP_API_KEY")
    base = "https://financialmodelingprep.com/api/v3"

    def get(endpoint):
        try:
            r = requests.get(f"{base}/{endpoint}/{ticker_bk}?limit=4&apikey={api_key}", timeout=15)
            r.raise_for_status()
            data = r.json()
            return data if isinstance(data, list) else []
        except Exception:
            return []

    income_data  = get("income-statement")
    balance_data = get("balance-sheet-statement")
    metrics_data = get("key-metrics")
    cashflow_data = get("cash-flow-statement")

    if not income_data:
        return {"error": f"No financial data found for {ticker_bk}"}

    years            = [item["date"][:4]                  for item in income_data[:4]]
    revenue          = [_fmt(item.get("revenue"))         for item in income_data[:4]]
    ebitda           = [_fmt(item.get("ebitda"))          for item in income_data[:4]]
    net_income       = [_fmt(item.get("netIncome"))       for item in income_data[:4]]
    operating_income = [_fmt(item.get("operatingIncome")) for item in income_data[:4]]

    total_debt   = [_fmt(item.get("totalDebt"))   for item in balance_data[:4]]
    total_assets = [_fmt(item.get("totalAssets")) for item in balance_data[:4]]
    total_equity = [_fmt(item.get("totalEquity")) for item in balance_data[:4]]

    m  = metrics_data[0]  if metrics_data  else {}
    i0 = income_data[0]   if income_data   else {}
    cf = cashflow_data[0] if cashflow_data else {}

    roe          = m.get("roe")
    ebitda_val   = float(i0.get("ebitda")  or 0)
    revenue_val  = float(i0.get("revenue") or 1)
    ebitda_margin = ebitda_val / revenue_val if revenue_val else None
    profit_margin = m.get("netProfitMargin")
    de = m.get("debtToEquity")

    return {
        "ticker_bk": ticker_bk,
        "years": years,
        "income": {
            "revenue":          revenue,
            "ebitda":           ebitda,
            "net_income":       net_income,
            "operating_income": operating_income,
        },
        "balance": {
            "total_debt":   total_debt,
            "total_assets": total_assets,
            "total_equity": total_equity,
        },
        "ratios": {
            "debt_to_equity": round(float(de), 2)              if de           is not None else None,
            "current_ratio":  m.get("currentRatio"),
            "roe":            round(float(roe) * 100, 1)       if roe          is not None else None,
            "ebitda_margin":  round(ebitda_margin * 100, 1)    if ebitda_margin             else None,
            "profit_margin":  round(float(profit_margin) * 100, 1) if profit_margin         else None,
            "free_cash_flow": _fmt(cf.get("freeCashFlow")),
        },
        "source": f"Financial Modeling Prep — {ticker_bk}",
    }


def fetch_company_research_ca(company: dict) -> str:
    name = company.get("english_name") or company.get("thai_name", "")
    ticker = company.get("ticker", "")
    search_name = f"{name} {ticker}".strip()

    results = []
    queries = [
        f"{search_name} company profile business overview Thailand",
        f"{search_name} news risk outlook 2024 2025",
        f"{search_name} industry Thailand market position",
    ]
    for q in queries:
        try:
            r = tavily.search(query=q, search_depth="basic", max_results=4, days=365)
            for item in r["results"]:
                results.append(f"[{item.get('title', '')}]\n{item['content']}\nURL: {item['url']}")
        except Exception:
            pass

    return "\n---\n".join(results[:10])


def draft_credit_memo(company: dict, financials: dict, research: str, facility: dict) -> dict:
    name = company.get("english_name") or company.get("thai_name", "")
    years = financials.get("years", [])
    income = financials.get("income", {})
    balance = financials.get("balance", {})
    ratios = financials.get("ratios", {})

    fin_lines = f"Financial Summary (M THB) | Years: {', '.join(years)}\n"
    fin_lines += f"Revenue:          {income.get('revenue')}\n"
    fin_lines += f"EBITDA:           {income.get('ebitda')}\n"
    fin_lines += f"Net Income:       {income.get('net_income')}\n"
    fin_lines += f"Total Debt:       {balance.get('total_debt')}\n"
    fin_lines += f"Total Assets:     {balance.get('total_assets')}\n"
    fin_lines += f"Total Equity:     {balance.get('total_equity')}\n"
    fin_lines += f"D/E: {ratios.get('debt_to_equity')}%  Current Ratio: {ratios.get('current_ratio')}x  ROE: {ratios.get('roe')}%  EBITDA Margin: {ratios.get('ebitda_margin')}%  FCF: {ratios.get('free_cash_flow')} M THB\n"

    amount = facility.get("amount", 0)
    collateral_val = facility.get("collateral_value", 0)
    ltv = f"{round(amount / collateral_val * 100, 1)}%" if collateral_val else "N/A"

    facility_lines = (
        f"Facility Type: {facility.get('facility_type')}\n"
        f"Amount: {amount / 1_000_000:,.0f} M THB\n"
        f"Tenor: {facility.get('tenor')} years\n"
        f"Collateral: {facility.get('collateral_type')} | Value: {collateral_val / 1_000_000:,.0f} M THB | LTV: {ltv}\n"
    )

    prompt = f"""You are a Thai bank credit analyst. Draft a credit application memo.

Company: {name} | Industry: {company.get("industry", "")}

{facility_lines}
{fin_lines}

Research:
{research[:3000]}

Draft these 4 sections in professional banking English. Return JSON only (no markdown):
{{
  "borrower_profile": "2-3 paragraphs: business description, ownership, market position, recent developments",
  "repayment_analysis": "primary repayment source, cash flow adequacy vs debt service, stress test at 20% revenue decline",
  "key_risks": [
    {{"risk": "...", "mitigant": "..."}},
    {{"risk": "...", "mitigant": "..."}},
    {{"risk": "...", "mitigant": "..."}}
  ],
  "recommendation": "Approve / Approve with conditions / Decline — with clear rationale and proposed conditions if any"
}}"""

    response = claude.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=2000,
        messages=[{"role": "user", "content": prompt}],
    )
    text = response.content[0].text.strip()
    text = re.sub(r"^```(?:json)?\s*", "", text)
    text = re.sub(r"\s*```$", "", text)
    try:
        return json.loads(text)
    except Exception:
        match = re.search(r"\{.*\}", text, re.DOTALL)
        if match:
            try:
                return json.loads(match.group())
            except Exception:
                pass
    return {"error": text}


def generate_ca_word_doc(company: dict, financials: dict, memo: dict, facility: dict) -> io.BytesIO:
    from docx import Document

    doc = Document()
    name = company.get("english_name") or company.get("thai_name", "")
    ticker = company.get("ticker", "—")

    doc.add_heading("CREDIT APPLICATION", 0)
    doc.add_heading(f"{name} ({ticker})", 1)

    # Facility summary
    doc.add_heading("Facility Summary", level=2)
    amount = facility.get("amount", 0)
    collateral_val = facility.get("collateral_value", 0)
    ltv = f"{round(amount / collateral_val * 100, 1)}%" if collateral_val else "N/A"

    fs = doc.add_table(rows=3, cols=4)
    fs.style = "Table Grid"
    fs.rows[0].cells[0].text = "Facility Type";  fs.rows[0].cells[1].text = facility.get("facility_type", "—")
    fs.rows[0].cells[2].text = "Amount (THB)";   fs.rows[0].cells[3].text = f"{amount:,.0f}"
    fs.rows[1].cells[0].text = "Tenor";          fs.rows[1].cells[1].text = f"{facility.get('tenor', '—')} years"
    fs.rows[1].cells[2].text = "Industry";       fs.rows[1].cells[3].text = company.get("industry", "—")
    fs.rows[2].cells[0].text = "Collateral";     fs.rows[2].cells[1].text = facility.get("collateral_type", "—")
    fs.rows[2].cells[2].text = "LTV";            fs.rows[2].cells[3].text = ltv
    doc.add_paragraph()

    # 1. Borrower Profile
    doc.add_heading("1. Borrower Profile", level=1)
    doc.add_paragraph(memo.get("borrower_profile", "—"))
    doc.add_paragraph()

    # 2. Financial Summary
    doc.add_heading("2. Financial Summary", level=1)
    years = financials.get("years", [])
    income = financials.get("income", {})
    balance = financials.get("balance", {})
    ratios = financials.get("ratios", {})

    if years:
        ft = doc.add_table(rows=1, cols=len(years) + 1)
        ft.style = "Table Grid"
        hdr = ft.rows[0].cells
        hdr[0].text = "M THB"
        for i, y in enumerate(years):
            hdr[i + 1].text = y

        for label, vals in [
            ("Revenue",       income.get("revenue", [])),
            ("EBITDA",        income.get("ebitda", [])),
            ("Net Income",    income.get("net_income", [])),
            ("Total Debt",    balance.get("total_debt", [])),
            ("Total Assets",  balance.get("total_assets", [])),
            ("Total Equity",  balance.get("total_equity", [])),
        ]:
            row = ft.add_row().cells
            row[0].text = label
            for i, v in enumerate(vals[: len(years)]):
                row[i + 1].text = f"{v:,.1f}" if v is not None else "—"

    doc.add_paragraph()
    rp = doc.add_paragraph()
    rp.add_run("Key Ratios: ").bold = True
    rp.add_run(
        f"D/E {ratios.get('debt_to_equity', '—')}%  |  "
        f"Current Ratio {ratios.get('current_ratio', '—')}x  |  "
        f"ROE {ratios.get('roe', '—')}%  |  "
        f"EBITDA Margin {ratios.get('ebitda_margin', '—')}%  |  "
        f"FCF {ratios.get('free_cash_flow', '—')} M THB"
    )
    sp = doc.add_paragraph()
    sp.add_run(f"Source: {financials.get('source', 'Yahoo Finance / SET')}").italic = True
    doc.add_paragraph()

    # 3. Facility Details
    doc.add_heading("3. Facility Details", level=1)
    purpose = "Working Capital" if "Working" in facility.get("facility_type", "") else "Capital Expenditure / General Corporate"
    for label, val in [
        ("Facility Type",     facility.get("facility_type", "—")),
        ("Requested Amount",  f"THB {amount / 1_000_000:,.0f}M"),
        ("Tenor",             f"{facility.get('tenor', '—')} years"),
        ("Purpose",           purpose),
        ("Collateral Type",   facility.get("collateral_type", "—")),
        ("Collateral Value",  f"THB {collateral_val / 1_000_000:,.0f}M"),
        ("LTV",               ltv),
    ]:
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(val)
    doc.add_paragraph()

    # 4. Repayment Analysis
    doc.add_heading("4. Repayment Analysis", level=1)
    doc.add_paragraph(memo.get("repayment_analysis", "—"))
    doc.add_paragraph()

    # 5. Key Risks
    doc.add_heading("5. Key Risks", level=1)
    risks = memo.get("key_risks", [])
    if risks:
        rt = doc.add_table(rows=1, cols=3)
        rt.style = "Table Grid"
        rt.rows[0].cells[0].text = "#"
        rt.rows[0].cells[1].text = "Risk"
        rt.rows[0].cells[2].text = "Mitigant"
        for i, r in enumerate(risks, 1):
            row = rt.add_row().cells
            row[0].text = str(i)
            row[1].text = r.get("risk", "")
            row[2].text = r.get("mitigant", "")
    doc.add_paragraph()

    # 6. RM Recommendation
    doc.add_heading("6. RM Recommendation", level=1)
    doc.add_paragraph(memo.get("recommendation", "—"))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
