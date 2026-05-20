import io
import os
import streamlit as st
from dotenv import load_dotenv
from researcher import resolve_company, fetch_company_research, analyze_esg_opportunities

load_dotenv()

st.set_page_config(
    page_title="ESG Opportunity Finder",
    page_icon="🌿",
    layout="wide",
)

st.title("🌿 ESG Opportunity Finder")
st.caption("For Bank RMs — identify ESG banking opportunities for Thai business clients")

with st.sidebar:
    st.header("How to use")
    st.markdown("""
1. Enter a company name (Thai/English) or SET ticker
2. Click **Analyze**
3. Review ESG evidence and product opportunities

**Example inputs:**
- `PTT` or `ปตท`
- `RATCH` or `ratch`
- `SCG` or `ปูนซิเมนต์ไทย`
- `GULF`
- `Banpu` or `บ้านปู`
- `ThaiBev`
- `Central Retail`
    """)
    st.divider()
    if not os.getenv("ANTHROPIC_API_KEY") or not os.getenv("TAVILY_API_KEY"):
        st.error("⚠️ API keys missing. Create a `.env` file from `.env.example`")


def generate_word_doc(company, result):
    from docx import Document
    from docx.shared import Pt

    doc = Document()

    label = f"{company.get('english_name', '')} ({company.get('ticker', 'non-listed')})"
    doc.add_heading(f"ESG Brief: {label}", 0)

    # Company profile table
    doc.add_heading("Company Profile", level=1)
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    cells = tbl.rows[0].cells
    cells[0].text = f"Industry: {company.get('industry', '—')}"
    cells[1].text = f"Listed: {'SET Listed' if company.get('listed') else 'Non-listed'}"
    cells[2].text = f"Ticker: {company.get('ticker', '—')}"
    doc.add_paragraph()

    # ESG Signals
    doc.add_heading("ESG Evidence & Signals", level=1)
    signals = result.get("esg_signals", [])
    if signals:
        for i, signal in enumerate(signals, 1):
            date = signal.get("date") or "Date unknown"
            related = signal.get("related_opportunity", "")
            p = doc.add_paragraph()
            run = p.add_run(f"{i}. [{date}]")
            run.bold = True
            if related:
                p.add_run(f"  —  {related}")
            doc.add_paragraph(signal.get("finding", ""))
            source_title = signal.get("source_title", "")
            source_url = signal.get("source_url", "")
            src_text = f"Source: {source_title}"
            if source_url:
                src_text += f" — {source_url}"
            doc.add_paragraph(src_text).runs[0].italic = True
            doc.add_paragraph()
    else:
        doc.add_paragraph("No ESG signals found.")

    doc.add_paragraph()

    # Opportunities
    doc.add_heading("ESG Banking Opportunities", level=1)
    opportunities = result.get("opportunities", [])
    if opportunities:
        for i, opp in enumerate(opportunities, 1):
            doc.add_heading(f"{i}. {opp.get('product', '')}", level=2)
            doc.add_paragraph(opp.get("rationale", ""))

            size_p = doc.add_paragraph()
            size_p.add_run("Estimated Deal Size: ").bold = True
            size_p.add_run(opp.get("size_thb", "—"))

            calc = opp.get("size_calculation", "")
            if calc:
                calc_p = doc.add_paragraph()
                calc_p.add_run("Size Calculation: ").bold = True
                calc_p.add_run(calc)

            assumptions = opp.get("assumptions", [])
            if assumptions:
                assump_p = doc.add_paragraph()
                assump_p.add_run("Key Assumptions:").bold = True
                for a in assumptions:
                    doc.add_paragraph(a, style="List Bullet")
            doc.add_paragraph()
    else:
        doc.add_paragraph("No opportunities identified.")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


tab1, tab2 = st.tabs(["ESG Finder", "ESG Finder + Word Doc"])

# ── TAB 1 (original, untouched) ─────────────────────────────────────────────
with tab1:
    if "analyses" not in st.session_state:
        st.session_state.analyses = []

    col1, col2 = st.columns([4, 1])
    with col1:
        company_input = st.text_input(
            "Company name or ticker",
            placeholder="e.g. PTT, RATCH, SCG, GULF, บ้านปู, Central Retail",
            label_visibility="collapsed",
        )
    with col2:
        analyze_btn = st.button("Analyze", type="primary", use_container_width=True)

    if analyze_btn and company_input:
        if not os.getenv("ANTHROPIC_API_KEY") or not os.getenv("TAVILY_API_KEY"):
            st.error("Please add your API keys to the `.env` file first.")
        else:
            with st.status(f"Researching **{company_input}**...", expanded=True) as status:
                st.write("🔍 Resolving company identity...")
                company = resolve_company(company_input)

                if "error" in company:
                    status.update(label="Company not found", state="error")
                    st.error(f"Could not find: **{company_input}**. Try a different name or ticker.")
                else:
                    st.write(f"✅ Found: **{company.get('english_name')}** ({company.get('ticker', 'non-listed')})")
                    st.write("📄 Fetching annual reports, sustainability reports, and news...")
                    research = fetch_company_research(company)

                    st.write("🤖 Analyzing ESG opportunities...")
                    result = analyze_esg_opportunities(company, research)

                    status.update(label="Analysis complete", state="complete", expanded=False)

                    st.session_state.analyses.insert(0, {
                        "query": company_input,
                        "company": company,
                        "result": result,
                    })

    for item in st.session_state.analyses:
        company = item["company"]
        result = item["result"]
        label = f"{company.get('english_name', item['query'])} ({company.get('ticker', 'non-listed')})"

        with st.expander(f"🏢 {label}", expanded=(item == st.session_state.analyses[0])):
            col_a, col_b, col_c = st.columns(3)
            col_a.metric("Industry", company.get("industry", "—"))
            col_b.metric("Listed", "SET Listed" if company.get("listed") else "Non-listed")
            col_c.metric("Ticker", company.get("ticker", "—"))

            if "error" in result:
                st.error("Could not parse analysis output.")
                st.code(result["error"])
                continue

            st.divider()

            st.subheader("📌 ESG Evidence & Signals")
            st.caption("Key facts, announcements, and commitments directly relevant to identified opportunities")

            signals = result.get("esg_signals", [])
            if signals:
                for i, signal in enumerate(signals, 1):
                    with st.container():
                        date = signal.get("date") or "Date unknown"
                        related = signal.get("related_opportunity", "")
                        url = signal.get("source_url", "")
                        title = signal.get("source_title", "Source")

                        col_left, col_right = st.columns([3, 1])
                        with col_left:
                            st.markdown(
                                f"**{i}. `{date}`** &nbsp; "
                                + (f"<span style='background:#e8f4ea;color:#2d6a4f;padding:2px 8px;border-radius:4px;font-size:0.8em'>{related}</span>" if related else ""),
                                unsafe_allow_html=True,
                            )
                            st.markdown(signal.get("finding", ""))
                        with col_right:
                            if url:
                                st.markdown(f"[🔗 {title}]({url})")
                            else:
                                st.markdown(f"📄 {title}")
                        st.divider()
            else:
                st.info("No ESG signals found.")

            st.subheader("💼 ESG Banking Opportunities")
            st.caption("Suggested products, deal size with calculation, and key assumptions")

            opportunities = result.get("opportunities", [])
            if opportunities:
                for i, opp in enumerate(opportunities, 1):
                    with st.container():
                        col_left, col_right = st.columns([3, 1])
                        with col_left:
                            st.markdown(f"### {i}. {opp.get('product', '')}")
                            st.markdown(opp.get("rationale", ""))
                            calc = opp.get("size_calculation", "")
                            if calc:
                                st.markdown("**Size Calculation:**")
                                st.info(calc)
                            assumptions = opp.get("assumptions", [])
                            if assumptions:
                                st.markdown("**Key Assumptions:**")
                                for a in assumptions:
                                    st.markdown(f"- {a}")
                        with col_right:
                            st.metric("Estimated Deal Size", opp.get("size_thb", "—"))
                        st.divider()
            else:
                st.info("No opportunities identified.")

            lines = [f"ESG Brief: {label}\n", "=" * 60]
            lines.append("\nESG EVIDENCE & SIGNALS\n")
            for s in signals:
                lines.append(f"• {s.get('finding', '')}")
                lines.append(f"  → {s.get('implication', '')}")
                if s.get("source_url"):
                    lines.append(f"  Source: {s.get('source_url')}")
            lines.append("\n\nESG BANKING OPPORTUNITIES\n")
            for o in opportunities:
                lines.append(f"• {o.get('product')} | Size: {o.get('size_thb', '—')}")
                lines.append(f"  {o.get('rationale', '')}")
                for a in o.get("assumptions", []):
                    lines.append(f"  - {a}")

            st.download_button(
                label="Download brief",
                data="\n".join(lines),
                file_name=f"esg_{company.get('ticker', item['query']).replace(' ', '_')}.txt",
                mime="text/plain",
                key=f"dl_t1_{item['query']}_{id(item)}",
            )

# ── TAB 2 (Word doc download) ────────────────────────────────────────────────
with tab2:
    if "analyses_word" not in st.session_state:
        st.session_state.analyses_word = []

    col1, col2 = st.columns([4, 1])
    with col1:
        company_input_w = st.text_input(
            "Company name or ticker",
            placeholder="e.g. PTT, RATCH, SCG, GULF, บ้านปู, Central Retail",
            label_visibility="collapsed",
            key="company_input_word",
        )
    with col2:
        analyze_btn_w = st.button("Analyze", type="primary", use_container_width=True, key="analyze_btn_word")

    if analyze_btn_w and company_input_w:
        if not os.getenv("ANTHROPIC_API_KEY") or not os.getenv("TAVILY_API_KEY"):
            st.error("Please add your API keys to the `.env` file first.")
        else:
            with st.status(f"Researching **{company_input_w}**...", expanded=True) as status:
                st.write("🔍 Resolving company identity...")
                company = resolve_company(company_input_w)

                if "error" in company:
                    status.update(label="Company not found", state="error")
                    st.error(f"Could not find: **{company_input_w}**. Try a different name or ticker.")
                else:
                    st.write(f"✅ Found: **{company.get('english_name')}** ({company.get('ticker', 'non-listed')})")
                    st.write("📄 Fetching annual reports, sustainability reports, and news...")
                    research = fetch_company_research(company)

                    st.write("🤖 Analyzing ESG opportunities...")
                    result = analyze_esg_opportunities(company, research)

                    status.update(label="Analysis complete", state="complete", expanded=False)

                    st.session_state.analyses_word.insert(0, {
                        "query": company_input_w,
                        "company": company,
                        "result": result,
                    })

    for item in st.session_state.analyses_word:
        company = item["company"]
        result = item["result"]
        label = f"{company.get('english_name', item['query'])} ({company.get('ticker', 'non-listed')})"

        with st.expander(f"🏢 {label}", expanded=(item == st.session_state.analyses_word[0])):
            col_a, col_b, col_c = st.columns(3)
            col_a.metric("Industry", company.get("industry", "—"))
            col_b.metric("Listed", "SET Listed" if company.get("listed") else "Non-listed")
            col_c.metric("Ticker", company.get("ticker", "—"))

            if "error" in result:
                st.error("Could not parse analysis output.")
                st.code(result["error"])
                continue

            st.divider()

            st.subheader("📌 ESG Evidence & Signals")
            st.caption("Key facts, announcements, and commitments directly relevant to identified opportunities")

            signals = result.get("esg_signals", [])
            if signals:
                for i, signal in enumerate(signals, 1):
                    with st.container():
                        date = signal.get("date") or "Date unknown"
                        related = signal.get("related_opportunity", "")
                        url = signal.get("source_url", "")
                        title = signal.get("source_title", "Source")

                        col_left, col_right = st.columns([3, 1])
                        with col_left:
                            st.markdown(
                                f"**{i}. `{date}`** &nbsp; "
                                + (f"<span style='background:#e8f4ea;color:#2d6a4f;padding:2px 8px;border-radius:4px;font-size:0.8em'>{related}</span>" if related else ""),
                                unsafe_allow_html=True,
                            )
                            st.markdown(signal.get("finding", ""))
                        with col_right:
                            if url:
                                st.markdown(f"[🔗 {title}]({url})")
                            else:
                                st.markdown(f"📄 {title}")
                        st.divider()
            else:
                st.info("No ESG signals found.")

            st.subheader("💼 ESG Banking Opportunities")
            st.caption("Suggested products, deal size with calculation, and key assumptions")

            opportunities = result.get("opportunities", [])
            if opportunities:
                for i, opp in enumerate(opportunities, 1):
                    with st.container():
                        col_left, col_right = st.columns([3, 1])
                        with col_left:
                            st.markdown(f"### {i}. {opp.get('product', '')}")
                            st.markdown(opp.get("rationale", ""))
                            calc = opp.get("size_calculation", "")
                            if calc:
                                st.markdown("**Size Calculation:**")
                                st.info(calc)
                            assumptions = opp.get("assumptions", [])
                            if assumptions:
                                st.markdown("**Key Assumptions:**")
                                for a in assumptions:
                                    st.markdown(f"- {a}")
                        with col_right:
                            st.metric("Estimated Deal Size", opp.get("size_thb", "—"))
                        st.divider()
            else:
                st.info("No opportunities identified.")

            word_buf = generate_word_doc(company, result)
            st.download_button(
                label="Download Word doc (.docx)",
                data=word_buf,
                file_name=f"esg_{company.get('ticker', item['query']).replace(' ', '_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key=f"dl_t2_{item['query']}_{id(item)}",
            )
